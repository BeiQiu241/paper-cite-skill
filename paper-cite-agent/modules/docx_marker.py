"""Write inline numeric citation marks back to a Word document."""

from __future__ import annotations

from typing import Any, Dict, List

from docx import Document


def _citation_label(cite_indices: List[int]) -> str:
    """Build compact inline citations such as [1] or [1][2]."""
    labels = [f"[{index + 1}]" for index in cite_indices if isinstance(index, int)]
    return "".join(labels) if labels else "[1]"


def annotate_docx(
    input_path: str,
    output_path: str,
    citation_positions: List[Dict[str, Any]],
    highlight_color: str = "yellow",
) -> str:
    """Append inline numeric citations to matched body paragraphs."""
    del highlight_color

    doc = Document(input_path)
    by_paragraph = {
        item["paragraph_index"]: item
        for item in citation_positions
        if isinstance(item, dict) and isinstance(item.get("paragraph_index"), int)
    }

    paragraph_index = 0
    inserted_total = 0

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue

        paragraph_index += 1
        position = by_paragraph.get(paragraph_index)
        if position is None:
            continue

        label = _citation_label(position.get("cite_indices", []))
        if label in paragraph.text:
            continue

        separator = "" if paragraph.text.endswith((" ", "\t")) else " "
        paragraph.add_run(f"{separator}{label}")
        inserted_total += 1

    doc.save(output_path)
    print(f"  Citation labels inserted: {inserted_total}")
    return output_path
