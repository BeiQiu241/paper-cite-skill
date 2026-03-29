"""参考文献生成模块：生成 GB/T 7714 / APA / IEEE 格式的参考文献列表。

GB/T 7714-2015（默认）：
  中文: [序号] 作者. 标题[J]. 期刊名, 年份, 卷(期): 页码. DOI.
  英文: [序号] LAST F, LAST F. Title[J]. Journal, Year, Vol(No): Pages. DOI.
APA：
  作者. (年). 标题. 期刊名, 卷(期), 页码.
IEEE：
  [序号] 作者, "标题," 期刊名, vol., no., pp., 年.
"""

import re
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt


def _is_zh_paper(paper: Dict[str, Any]) -> bool:
    """判断是否应使用中文格式（中文标题或明确标记为 zh）。"""
    if paper.get("lang") == "zh":
        return True
    title = paper.get("title", "")
    return bool(re.search(r'[\u4e00-\u9fff]', title))


def _cn_en_counts(papers: List[Dict[str, Any]]) -> tuple[int, int]:
    cn = sum(1 for p in papers if _is_zh_paper(p))
    en = len(papers) - cn
    return cn, en


def _journal_suffix_apa(paper: Dict[str, Any]) -> str:
    """拼接期刊信息后缀（APA 英文格式）。"""
    journal = paper.get("journal", "")
    volume = str(paper.get("volume", "") or "")
    issue = str(paper.get("issue", "") or "")
    pages = str(paper.get("pages", "") or "")

    if not journal:
        return ""
    parts = [f" *{journal}*"]
    if volume:
        parts.append(f", *{volume}*")
        if issue:
            parts.append(f"({issue})")
    if pages:
        parts.append(f", {pages}")
    return "".join(parts) + "."


def _journal_suffix_zh(paper: Dict[str, Any]) -> str:
    """拼接期刊信息后缀（中文格式）。"""
    journal = paper.get("journal", "")
    volume = str(paper.get("volume", "") or "")
    issue = str(paper.get("issue", "") or "")
    pages = str(paper.get("pages", "") or "")

    if not journal:
        return ""
    parts = [journal]
    if volume:
        parts.append(f", {volume}")
        if issue:
            parts.append(f"({issue})")
    if pages:
        parts.append(f", {pages}")
    return "".join(parts) + "."


def format_apa(paper: Dict[str, Any], index: int) -> str:
    """生成 APA 格式引用；中文论文自动使用中文学术格式。"""
    if _is_zh_paper(paper):
        return _format_zh(paper)

    authors = paper.get("authors", [])
    year = paper.get("year", "n.d.")
    title = paper.get("title", "Unknown Title")
    doi = paper.get("doi", "")
    url = paper.get("url", "")

    formatted_authors = []
    for author in authors[:6]:
        parts = author.strip().split()
        if len(parts) >= 2:
            last = parts[-1]
            initials = ". ".join(p[0] for p in parts[:-1] if p) + "."
            formatted_authors.append(f"{last}, {initials}")
        else:
            formatted_authors.append(author)

    if len(authors) > 6:
        author_str = ", ".join(formatted_authors[:6]) + ", ... " + formatted_authors[-1]
    elif len(formatted_authors) > 1:
        author_str = ", ".join(formatted_authors[:-1]) + ", & " + formatted_authors[-1]
    elif formatted_authors:
        author_str = formatted_authors[0]
    else:
        author_str = "Unknown Author"

    ref = f"{author_str} ({year}). {title}."

    journal_part = _journal_suffix_apa(paper)
    if journal_part:
        ref += " " + journal_part
    elif doi:
        ref += f" https://doi.org/{doi}"
    elif url:
        ref += f" {url}"

    return ref


def _format_zh(paper: Dict[str, Any]) -> str:
    """中文学术格式：作者. (年). 标题. 期刊名, 卷(期), 页码."""
    authors = paper.get("authors", [])
    year = paper.get("year", "n.d.")
    title = paper.get("title", "")
    doi = paper.get("doi", "")
    url = paper.get("url", "")

    # 中文作者直接拼接，英文名保持原样
    author_str = ", ".join(authors[:6])
    if len(authors) > 6:
        author_str += ", 等"
    if not author_str:
        author_str = "佚名"

    ref = f"{author_str}. ({year}). {title}."

    journal_part = _journal_suffix_zh(paper)
    if journal_part:
        ref += " " + journal_part
    elif doi:
        ref += f" https://doi.org/{doi}"
    elif url:
        ref += f" {url}"

    return ref


def format_ieee(paper: Dict[str, Any], index: int) -> str:
    """生成 IEEE 格式引用；中文论文保留期刊信息。"""
    authors = paper.get("authors", [])
    year = paper.get("year", "n.d.")
    title = paper.get("title", "Unknown Title")
    doi = paper.get("doi", "")
    url = paper.get("url", "")
    journal = paper.get("journal", "")
    volume = str(paper.get("volume", "") or "")
    issue = str(paper.get("issue", "") or "")
    pages = str(paper.get("pages", "") or "")

    formatted_authors = []
    for author in authors[:6]:
        parts = author.strip().split()
        if len(parts) >= 2:
            last = parts[-1]
            initials = ". ".join(p[0] for p in parts[:-1] if p) + "."
            formatted_authors.append(f"{initials} {last}")
        else:
            formatted_authors.append(author)

    if len(authors) > 6:
        author_str = ", ".join(formatted_authors[:6]) + " et al."
    else:
        author_str = ", ".join(formatted_authors)

    ref = f"[{index}] {author_str}, \"{title},\""
    if journal:
        ref += f" {journal},"
    if volume:
        ref += f" vol. {volume},"
        if issue:
            ref += f" no. {issue},"
    if pages:
        ref += f" pp. {pages},"
    ref += f" {year}."
    if doi:
        ref += f" doi: {doi}."
    elif url:
        ref += f" [Online]. Available: {url}"

    return ref


def _format_gbt7714_authors_zh(authors: List[str]) -> str:
    """中文 GB/T 7714 作者格式：直接拼接，超6人加"等"。"""
    author_str = ", ".join(authors[:6])
    if len(authors) > 6:
        author_str += ", 等"
    return author_str or "佚名"


def _format_gbt7714_authors_en(authors: List[str]) -> str:
    """英文 GB/T 7714 作者格式：LAST F（大写），超6人加 et al。"""
    formatted = []
    for author in authors[:6]:
        parts = author.strip().split()
        if len(parts) >= 2:
            last = parts[-1].upper()
            initials = " ".join(p[0].upper() + "." for p in parts[:-1] if p)
            formatted.append(f"{last} {initials}")
        else:
            formatted.append(author.upper())
    if len(authors) > 6:
        return ", ".join(formatted) + ", et al"
    return ", ".join(formatted) or "Unknown Author"


def _gbt7714_journal_part(paper: Dict[str, Any]) -> str:
    """拼接 GB/T 7714 期刊/卷期/页码部分。"""
    journal = paper.get("journal", "")
    volume = str(paper.get("volume", "") or "")
    issue = str(paper.get("issue", "") or "")
    pages = str(paper.get("pages", "") or "")
    year = str(paper.get("year", "") or "")

    parts = []
    if journal:
        parts.append(f" {journal},")
    if year:
        parts.append(f" {year}")
    if volume:
        parts.append(f", {volume}")
        if issue:
            parts.append(f"({issue})")
    if pages:
        parts.append(f": {pages}")
    return "".join(parts)


def format_gbt7714(paper: Dict[str, Any], index: int) -> str:
    """生成 GB/T 7714-2015 格式引用（期刊文章 [J]）。"""
    authors = paper.get("authors", [])
    title = paper.get("title", "")
    doi = paper.get("doi", "")
    url = paper.get("url", "")

    if _is_zh_paper(paper):
        author_str = _format_gbt7714_authors_zh(authors)
    else:
        author_str = _format_gbt7714_authors_en(authors)

    ref = f"[{index}] {author_str}. {title}[J]."
    ref += _gbt7714_journal_part(paper)
    ref += "."

    if doi:
        ref += f" DOI: {doi}."
    elif url:
        ref += f" {url}"

    return ref


def generate_reference_list(
    papers: List[Dict[str, Any]],
    fmt: str = "GBT7714",
) -> List[str]:
    """生成完整参考文献列表。"""
    references = []
    fmt_upper = fmt.upper().replace("-", "").replace(" ", "").replace("/", "")
    for i, paper in enumerate(papers, start=1):
        if fmt_upper == "IEEE":
            ref = format_ieee(paper, i)
        elif fmt_upper in ("APA",):
            ref = format_apa(paper, i)
        else:  # GBT7714 / GB/T7714 / default
            ref = format_gbt7714(paper, i)
        references.append(ref)
    return references


def save_references_txt(
    references: List[str],
    output_path: str,
    all_candidates: List[Dict[str, Any]] = None,
    selected_papers: List[Dict[str, Any]] = None,
    fmt: str = "APA",
):
    """
    保存参考文献列表到 txt 文件。

    - references      : 推荐文献（已格式化字符串，写入"推荐参考文献"区块）
    - all_candidates  : 全部候选文献原始列表，写入"全部候选文献"区块
    - fmt             : 全部候选的格式化格式
    """
    with open(output_path, "w", encoding="utf-8") as f:
        # ── 推荐文献 ──────────────────────────────────────────────
        if selected_papers is not None:
            cn, en = _cn_en_counts(selected_papers)
            f.write(f"[LLM审查] 选出 {len(selected_papers)} 篇（中文 {cn} 篇 / 英文 {en} 篇）\n")
        else:
            f.write(f"[LLM审查] 选出 {len(references)} 篇\n")
        f.write("=" * 60 + "\n\n")
        for ref in references:
            f.write(ref + "\n\n")

        # ── 全部候选 ──────────────────────────────────────────────
        if all_candidates:
            f.write("\n\n【全部候选文献（按相关度排序）】\n")
            f.write("=" * 60 + "\n")
            f.write(f"共 {len(all_candidates)} 篇\n\n")

            cn_papers = [p for p in all_candidates
                         if p.get("source", "").endswith("_cn") or p.get("lang") == "zh"]
            en_papers = [p for p in all_candidates
                         if not (p.get("source", "").endswith("_cn") or p.get("lang") == "zh")]

            if cn_papers:
                f.write(f"── 中国机构/中文文献（{len(cn_papers)} 篇）──\n\n")
                for i, paper in enumerate(cn_papers, 1):
                    ref = format_apa(paper, i) if fmt.upper() != "IEEE" else format_ieee(paper, i)
                    score = paper.get("_score")
                    score_str = f"  [相关度: {score:.3f}]" if score is not None else ""
                    f.write(f"{ref}{score_str}\n\n")

            if en_papers:
                f.write(f"── 英文文献（{len(en_papers)} 篇）──\n\n")
                for i, paper in enumerate(en_papers, 1):
                    ref = format_apa(paper, i) if fmt.upper() != "IEEE" else format_ieee(paper, i)
                    score = paper.get("_score")
                    score_str = f"  [相关度: {score:.3f}]" if score is not None else ""
                    f.write(f"{ref}{score_str}\n\n")

    print(f"  [References] 参考文献已保存至: {output_path}")


def append_references_to_docx(
    doc_path: str,
    output_path: str,
    references: List[str],
    papers: List[Dict[str, Any]],
):
    """将参考文献列表追加写入 Word 文档末尾。"""
    doc = Document(doc_path)

    doc.add_page_break()
    cn, en = _cn_en_counts(papers)
    heading = doc.add_heading(
        f"[LLM审查] 选出 {len(papers)} 篇（中文 {cn} 篇 / 英文 {en} 篇）",
        level=1,
    )

    for i, (ref, paper) in enumerate(zip(references, papers), start=1):
        para = doc.add_paragraph(ref)
        para.style = doc.styles["Normal"]

        # 添加链接信息
        url = paper.get("url", "")
        pdf_url = paper.get("pdf_url", "")
        if url or pdf_url:
            link_text = []
            if url:
                link_text.append(f"链接: {url}")
            if pdf_url:
                link_text.append(f"PDF: {pdf_url}")
            note_para = doc.add_paragraph("    " + " | ".join(link_text))
            note_para.runs[0].font.size = Pt(9)

        doc.add_paragraph()  # 空行分隔

    doc.save(output_path)
    print(f"  [References] 参考文献已写入文档: {output_path}")
