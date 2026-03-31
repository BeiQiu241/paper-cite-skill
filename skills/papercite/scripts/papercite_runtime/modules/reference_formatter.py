"""Reference formatting helpers."""

from __future__ import annotations

import re
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ZH_ET_AL = "\u7b49"
ZH_REFERENCE_HEADING = "\u53c2\u8003\u6587\u732e"
ZH_SONGTI = "\u5b8b\u4f53"
ZH_HEITI = "\u9ed1\u4f53"


def _is_zh_paper(paper: Dict[str, Any]) -> bool:
    """Infer whether a paper should be formatted as Chinese."""
    return str(paper.get("lang", "")).lower() == "zh" or bool(
        re.search(r"[\u4e00-\u9fff]", str(paper.get("title", "")))
    )


def _author_list(paper: Dict[str, Any]) -> List[str]:
    """Return cleaned authors."""
    authors = paper.get("authors") or []
    cleaned = [str(author).strip() for author in authors if str(author).strip()]
    return cleaned


def _join_authors(paper: Dict[str, Any]) -> str:
    """Join authors in a compact GB/T-style list."""
    authors = _author_list(paper)
    if not authors:
        return "Unknown Author"
    if len(authors) > 3:
        shown = authors[:3]
        suffix = ZH_ET_AL if _is_zh_paper(paper) else "et al"
        return f"{', '.join(shown)}, {suffix}"
    return ", ".join(authors)


def _source_parts(paper: Dict[str, Any]) -> str:
    """Build journal, year, volume, issue, and pages text."""
    journal = str(paper.get("journal") or "").strip()
    year = str(paper.get("year") or "").strip()
    volume = str(paper.get("volume") or "").strip()
    issue = str(paper.get("issue") or "").strip()
    pages = str(paper.get("pages") or "").strip()

    parts: List[str] = []
    if journal:
        parts.append(journal)

    year_part = year
    if volume:
        year_part = f"{year_part}, {volume}" if year_part else volume
        if issue:
            year_part = f"{year_part}({issue})"
    elif issue:
        year_part = f"{year_part}({issue})" if year_part else f"({issue})"

    if pages:
        year_part = f"{year_part}: {pages}" if year_part else pages

    if year_part:
        parts.append(year_part)
    return ". ".join(parts)


def _append_locator(ref: str, paper: Dict[str, Any]) -> str:
    """Append DOI or URL when present."""
    doi = str(paper.get("doi") or "").strip()
    url = str(paper.get("url") or "").strip()
    if doi:
        return f"{ref} DOI: {doi}."
    if url:
        return f"{ref} Available: {url}."
    return ref


def format_gbt7714(paper: Dict[str, Any], index: int) -> str:
    """Format one reference in simplified GB/T 7714-2015 journal style."""
    ref = f"[{index}] {_join_authors(paper)}. {paper.get('title', 'Untitled')}[J]."
    source = _source_parts(paper)
    if source:
        ref = f"{ref} {source}."
    return _append_locator(ref, paper)


def format_apa(paper: Dict[str, Any], index: int) -> str:
    """Format one reference in simplified APA style."""
    year = paper.get("year") or "n.d."
    ref = f"{_join_authors(paper)} ({year}). {paper.get('title', 'Untitled')}."
    source = _source_parts(paper)
    if source:
        ref = f"{ref} {source}."
    return _append_locator(ref, paper)


def format_ieee(paper: Dict[str, Any], index: int) -> str:
    """Format one reference in simplified IEEE style."""
    ref = f"[{index}] {_join_authors(paper)}, \"{paper.get('title', 'Untitled')}.\""
    source = _source_parts(paper)
    if source:
        ref = f"{ref} {source}."
    return _append_locator(ref, paper)


def generate_reference_list(
    papers: List[Dict[str, Any]],
    fmt: str = "GBT7714",
) -> List[str]:
    """Generate formatted references."""
    formatter = {
        "GBT7714": format_gbt7714,
        "GBT7714-2015": format_gbt7714,
        "GB/T7714": format_gbt7714,
        "GB/T7714-2015": format_gbt7714,
        "APA": format_apa,
        "IEEE": format_ieee,
    }.get(fmt.upper(), format_gbt7714)

    return [formatter(paper, index) for index, paper in enumerate(papers, start=1)]


def _numbered_reference(ref: str, index: int) -> str:
    """Keep numbering stable even for non-numbered styles."""
    if ref.lstrip().startswith("["):
        return ref
    return f"[{index}] {ref}"


def save_references_txt(references: List[str], output_path: str) -> None:
    """Write references to a plain text file."""
    with open(output_path, "w", encoding="utf-8") as handle:
        handle.write(f"{ZH_REFERENCE_HEADING}\n")
        handle.write("=" * 40 + "\n\n")
        for index, reference in enumerate(references, start=1):
            handle.write(_numbered_reference(reference, index) + "\n")
    print(f"  References saved: {output_path}")


def _apply_reference_run_style(run) -> None:
    """Apply mixed Chinese/English font settings required by the template."""
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), ZH_SONGTI)
    run.font.size = Pt(10.5)


def _add_reference_heading(doc: Document) -> None:
    """Append the formatted reference heading."""
    paragraph = doc.add_paragraph()
    try:
        paragraph.style = "Heading 1"
    except KeyError:
        pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(14)

    run = paragraph.add_run(ZH_REFERENCE_HEADING)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), ZH_HEITI)


def _add_reference_paragraph(doc: Document, text: str) -> None:
    """Append one formatted reference paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(-15.75)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(text)
    _apply_reference_run_style(run)


def append_references_to_docx(
    doc_path: str,
    output_path: str,
    references: List[str],
) -> None:
    """Append a formatted reference section to the end of a Word document."""
    doc = Document(doc_path)
    if references:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        _add_reference_heading(doc)
        for index, reference in enumerate(references, start=1):
            _add_reference_paragraph(doc, _numbered_reference(reference, index))

    doc.save(output_path)
    print(f"  References appended: {output_path}")
